from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOC = ROOT / "output" / "doc"
OUTPUT_PDF = ROOT / "output" / "pdf"
OUTPUT_PRESENTATION = ROOT / "output" / "presentation"

SLIDE_PDF = OUTPUT_PDF / "na-ryazan_presentation_backup.pdf"
NOTES_DOCX = OUTPUT_DOC / "na-ryazan_speaker_notes.docx"
NOTES_PDF = OUTPUT_PDF / "na-ryazan_speaker_notes.pdf"
OUTLINE_MD = OUTPUT_PRESENTATION / "na-ryazan_presentation_outline.md"

PAGE_WIDTH = 960
PAGE_HEIGHT = 540

BG = colors.HexColor("#12161F")
PANEL = colors.HexColor("#1A2231")
PANEL_ALT = colors.HexColor("#EEF3FB")
TEXT = colors.HexColor("#F4F7FB")
MUTED = colors.HexColor("#B8C6DD")
BLUE = colors.HexColor("#1E5BFF")
BLUE_SOFT = colors.HexColor("#7FAEFF")
GREEN = colors.HexColor("#8FF7B0")
RED = colors.HexColor("#FF6A6A")

IMAGES = {
    "hero": ROOT / "frontend" / "public" / "images" / "hero" / "00_ryazan.jpg",
    "kremlin": ROOT / "frontend" / "public" / "images" / "landmarks" / "ryazan-kremlin.jpg",
    "sunset": ROOT / "frontend" / "public" / "images" / "landmarks" / "kremlin-sunset.jpeg",
    "pryanik": ROOT / "frontend" / "public" / "images" / "landmarks" / "pryanik-museum.jpg",
    "riamz": ROOT / "frontend" / "public" / "images" / "landmarks" / "riamz.jpg",
    "vdv": ROOT / "frontend" / "public" / "images" / "landmarks" / "vdv-museum.jpg",
    "youth": ROOT / "frontend" / "public" / "images" / "landmarks" / "youth-house.jpg",
}

FONT_REGULAR = "Arial"
FONT_BOLD = "Arial-Bold"


SLIDES = [
    {
        "title": "На Рязань",
        "subtitle": "Интерактивный PWA-квест по Рязани",
        "tag": "Туризм / Культура / Городские маршруты",
        "bullets": [
            "Маршруты, QR-точки и геймификация в одном мобильном сервисе",
            "AI-гид, офлайн-режим, карта 2GIS и система наград",
            "Продукт превращает прогулку по городу в управляемый цифровой сценарий",
        ],
        "notes": [
            "Здравствуйте. Мы представляем проект На Рязань — интерактивный городской сервис, который делает знакомство с Рязанью не пассивной прогулкой, а увлекательным маршрутом с заданиями, контентом и наградами.",
            "Наш фокус — дать туристу готовый понятный сценарий, а городу и партнерам — цифровой инструмент вовлечения.",
            "Сегодня покажем проблему, решение, реализованный MVP, ценность для партнеров и roadmap развития.",
        ],
    },
    {
        "title": "Проблема",
        "subtitle": "Почему обычной прогулки уже недостаточно",
        "tag": "Контекст",
        "bullets": [
            "Туристу не хватает цельного цифрового сценария: что смотреть, куда идти и зачем задерживаться у точки",
            "Городские маршруты часто не удерживают внимание и не создают завершенный опыт",
            "Музеям и площадкам нужен дополнительный поток и понятная механика цифрового контакта с посетителем",
            "Нужен продукт, который соединяет навигацию, контент, игру и партнерскую выгоду",
        ],
        "notes": [
            "Проблема в том, что даже интересный город сам по себе не дает человеку удобного сценария движения.",
            "Часто турист открывает карту, хаотично выбирает точки и теряет интерес уже после первой остановки.",
            "Мы увидели возможность собрать целостный опыт: маршрут, контент, игру и партнерские механики внутри одного мобильного продукта.",
        ],
    },
    {
        "title": "Решение",
        "subtitle": "Что такое На Рязань",
        "tag": "Описание продукта",
        "bullets": [
            "PWA-сервис для путешествий по Рязани с мобильным сценарием использования",
            "Пользователь выбирает маршрут, проходит точки, сканирует QR-коды и получает баллы",
            "На каждой точке доступны факты, мини-квизы, элементы AI-гида и контент сопровождения",
            "Формат работает как культурный сервис, городской квест и партнерская платформа одновременно",
        ],
        "notes": [
            "На Рязань — это не просто сайт с описанием достопримечательностей.",
            "Это продукт, который сопровождает пользователя на прогулке и постоянно поддерживает интерес: через задания, знания, достижения и полезные остановки.",
            "За счет PWA-подхода сервис остается легким для запуска и удобным на телефоне.",
        ],
    },
    {
        "title": "Сценарий пользователя",
        "subtitle": "Как проходит один маршрут",
        "tag": "User Flow",
        "bullets": [
            "1. Пользователь выбирает маршрут по теме: история, популярные места, парки, вода",
            "2. Получает путь на карте и может добавить удобные остановки по дороге",
            "3. На точках сканирует QR-коды, открывает факты, аудио и игровые задания",
            "4. Завершает маршрут, проходит финальный тест и использует награды в партнерском контуре",
        ],
        "notes": [
            "Здесь важно показать, что путь пользователя очень простой и не требует отдельного обучения.",
            "Мы специально сделали механику знакомой: выбор маршрута, прохождение, подтверждение точки, награда.",
            "При этом внутри этого простого пути уже зашиты карта, контент, геймификация и партнерские касания.",
        ],
    },
    {
        "title": "MVP уже собран",
        "subtitle": "Что реализовано в продукте сейчас",
        "tag": "Функциональность",
        "bullets": [
            "4 маршрута, 11 точек интереса, 12 дополнительных остановок, 2 наградных сценария",
            "Карта маршрута, предпросмотр пути, избранные маршруты и профиль пользователя",
            "QR-сканер, офлайн-очередь сканов, лидерборд, магазин наград и финальные тесты",
            "AI-гид по точкам, карточки фактов и аудиосопровождение для части маршрута",
        ],
        "notes": [
            "Для нас важно подчеркнуть: это уже не идея и не просто дизайн-концепт.",
            "В репозитории реализованы реальные маршруты, логика точек, профиль, офлайн-хранение, сканирование и партнерский слой наград.",
            "То есть перед нами рабочий MVP, который можно тестировать, насыщать контентом и масштабировать.",
        ],
    },
    {
        "title": "Почему продукт удерживает",
        "subtitle": "Механика вовлечения пользователя",
        "tag": "Ценность для человека",
        "bullets": [
            "Город воспринимается как игра: каждая точка дает прогресс, баллы и новую информацию",
            "AI-гид, факты и квизы создают ощущение персонального сопровождения",
            "Награды и партнерские предложения добавляют практическую мотивацию пройти маршрут до конца",
            "Офлайн-режим снижает барьер использования прямо на прогулке",
        ],
        "notes": [
            "Удержание строится не на одном факторе, а на комбинации нескольких механик.",
            "Пользователь не только идет от точки к точке, но и постоянно получает подтверждение прогресса: баллы, задания, знания, бонусы.",
            "Это делает продукт интересным и для первого визита, и для повторного возвращения в сервис.",
        ],
    },
    {
        "title": "Ценность для города и партнеров",
        "subtitle": "Кому и зачем выгоден проект",
        "tag": "B2G / B2B",
        "bullets": [
            "Городу: современный цифровой слой поверх туристической среды и рост вовлеченности гостей",
            "Музеям и площадкам: дополнительный поток, QR-активации и цифровое сопровождение посещения",
            "Партнерам: промокоды, награды, билеты, спецпредложения и будущие платные сценарии",
            "Продукт можно развивать как городскую платформу взаимодействия, а не как разовую экскурсию",
        ],
        "notes": [
            "На Рязань интересен не только пользователю, но и экосистеме вокруг него.",
            "Город получает современный формат знакомства с локациями, музеи — дополнительное касание с посетителем, партнеры — конверсию в действие.",
            "Поэтому проект устойчив и как культурный сервис, и как платформа для сотрудничества.",
        ],
    },
    {
        "title": "Техническая архитектура",
        "subtitle": "На чем построен продукт",
        "tag": "Стек и реализация",
        "bullets": [
            "Frontend: React + Tailwind, PWA-подход и мобильный сценарий",
            "Backend: Node.js + Express с локальной JSON-базой для прототипного контура",
            "Интеграции: 2GIS для карты и маршрутизации, QR-сканирование и офлайн-хранилище",
            "Архитектура разделяет контент, маршрутную логику, пользовательский прогресс и награды",
        ],
        "notes": [
            "С технической точки зрения проект уже разбит на понятные модули: интерфейс, API, данные, контент и внешние интеграции.",
            "Это удобно для дальнейшего роста: можно отдельно масштабировать контент, партнерский слой, профиль или маршрутизацию.",
            "То есть решение не только демонстрационное, но и инженерно подготовлено к расширению.",
        ],
    },
    {
        "title": "Roadmap",
        "subtitle": "Как развивать На Рязань дальше",
        "tag": "Следующие шаги",
        "bullets": [
            "Этап 1. Пилот по Рязани: наполнение контентом, проверка сценариев, подключение партнерских наград",
            "Этап 2. Масштабирование по городу: новые маршруты, билеты, расширенный AI-гид и события",
            "Этап 3. Тиражирование механики на другие города и тематические культурные программы",
            "На Рязань — это база для современной городской туристической платформы",
        ],
        "notes": [
            "В ближайшей перспективе мы видим проект как сильный городской пилот по Рязани.",
            "Следующий шаг — расширять партнерскую экосистему, контентные сценарии и форматы монетизации.",
            "В долгую это не просто маршрутный сервис, а масштабируемая механика цифрового туризма, которую можно переносить на другие города и события.",
        ],
    },
]


def ensure_dirs() -> None:
    OUTPUT_DOC.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF.mkdir(parents=True, exist_ok=True)
    OUTPUT_PRESENTATION.mkdir(parents=True, exist_ok=True)


def register_fonts() -> None:
    global FONT_REGULAR, FONT_BOLD
    candidates = [
        ("C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/arialbd.ttf", "ArialCustom", "ArialCustom-Bold"),
        ("C:/Windows/Fonts/segoeui.ttf", "C:/Windows/Fonts/segoeuib.ttf", "SegoeUICustom", "SegoeUICustom-Bold"),
        ("C:/Windows/Fonts/calibri.ttf", "C:/Windows/Fonts/calibrib.ttf", "CalibriCustom", "CalibriCustom-Bold"),
    ]
    for regular_path, bold_path, regular_name, bold_name in candidates:
        if Path(regular_path).exists() and Path(bold_path).exists():
            pdfmetrics.registerFont(TTFont(regular_name, regular_path))
            pdfmetrics.registerFont(TTFont(bold_name, bold_path))
            FONT_REGULAR = regular_name
            FONT_BOLD = bold_name
            return


def draw_text_block(c: canvas.Canvas, lines: list[str], x: float, y: float, size: int = 14, leading: int = 19, color=TEXT, bold: bool = False) -> None:
    text = c.beginText()
    text.setTextOrigin(x, y)
    text.setFont(FONT_BOLD if bold else FONT_REGULAR, size)
    text.setLeading(leading)
    text.setFillColor(color)
    for line in lines:
        text.textLine(line)
    c.drawText(text)


def wrap_lines(text: str, width: int) -> list[str]:
    return wrap(text, width=width, break_long_words=False, break_on_hyphens=False)


def draw_footer(c: canvas.Canvas, index: int) -> None:
    c.setFillColor(BLUE)
    c.rect(0, 0, PAGE_WIDTH, 10, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont(FONT_REGULAR, 11)
    c.drawRightString(PAGE_WIDTH - 40, 24, f"{index + 1:02d}")
    c.drawString(40, 24, "На Рязань")


def draw_cover(c: canvas.Canvas, slide: dict, index: int) -> None:
    c.setFillColor(BG)
    c.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    image = ImageReader(str(IMAGES["hero"]))
    c.drawImage(image, 520, 0, width=440, height=540, mask="auto")
    c.setFillColor(colors.Color(0, 0, 0, alpha=0.38))
    c.rect(500, 0, 460, 540, fill=1, stroke=0)
    c.setFillColor(BLUE)
    c.rect(0, PAGE_HEIGHT - 24, PAGE_WIDTH, 24, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont(FONT_REGULAR, 14)
    c.drawString(42, PAGE_HEIGHT - 17, slide["tag"])
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 40)
    c.drawString(54, 390, slide["title"])
    c.setFont(FONT_REGULAR, 24)
    c.drawString(54, 352, slide["subtitle"])
    bullet_y = 292
    for bullet in slide["bullets"]:
        lines = wrap_lines(bullet, 44)
        c.setFillColor(BLUE_SOFT)
        c.circle(60, bullet_y + 6, 3.5, fill=1, stroke=0)
        draw_text_block(c, lines, 74, bullet_y, size=14, leading=18, color=MUTED)
        bullet_y -= 22 * len(lines) + 12
    c.setFillColor(colors.Color(1, 1, 1, alpha=0.14))
    c.roundRect(54, 56, 320, 76, 18, fill=1, stroke=0)
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 15)
    c.drawString(76, 103, "Ключевая идея")
    draw_text_block(
        c,
        wrap_lines("Собрать современный культурный маршрут как цифровой пользовательский продукт.", 38),
        76,
        80,
        size=13,
        leading=16,
        color=MUTED,
    )
    draw_footer(c, index)


def draw_standard_slide(c: canvas.Canvas, slide: dict, index: int, image_key: str | None = None, panel_style: str = "right") -> None:
    c.setFillColor(BG)
    c.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    c.setFillColor(colors.Color(1, 1, 1, alpha=0.05))
    c.roundRect(42, 56, 876, 430, 26, fill=1, stroke=0)
    c.setFillColor(BLUE)
    c.rect(42, 460, 210, 10, fill=1, stroke=0)
    c.setFillColor(MUTED)
    c.setFont(FONT_REGULAR, 14)
    c.drawString(42, 486, slide["tag"])
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 30)
    c.drawString(42, 430, slide["title"])
    c.setFont(FONT_REGULAR, 18)
    c.drawString(42, 402, slide["subtitle"])

    if image_key and panel_style == "right":
        c.setFillColor(PANEL_ALT)
        c.roundRect(560, 120, 320, 240, 24, fill=1, stroke=0)
        image = ImageReader(str(IMAGES[image_key]))
        c.drawImage(image, 578, 138, width=284, height=204, preserveAspectRatio=True, mask="auto")
    elif image_key and panel_style == "bottom":
        c.setFillColor(PANEL_ALT)
        c.roundRect(54, 86, 852, 160, 24, fill=1, stroke=0)
        image = ImageReader(str(IMAGES[image_key]))
        c.drawImage(image, 570, 102, width=316, height=128, preserveAspectRatio=True, mask="auto")

    bullet_x = 64
    bullet_y = 344
    bullet_width = 54 if image_key and panel_style == "right" else 76
    for bullet in slide["bullets"]:
        lines = wrap_lines(bullet, bullet_width)
        c.setFillColor(BLUE_SOFT)
        c.circle(bullet_x, bullet_y + 6, 3.2, fill=1, stroke=0)
        draw_text_block(c, lines, bullet_x + 14, bullet_y, size=15, leading=19, color=TEXT)
        bullet_y -= 22 * len(lines) + 12
    draw_footer(c, index)


def draw_process_slide(c: canvas.Canvas, slide: dict, index: int) -> None:
    c.setFillColor(BG)
    c.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 30)
    c.drawString(42, 454, slide["title"])
    c.setFont(FONT_REGULAR, 18)
    c.drawString(42, 426, slide["subtitle"])
    c.setFillColor(MUTED)
    c.setFont(FONT_REGULAR, 14)
    c.drawString(42, 482, slide["tag"])

    x_positions = [58, 286, 514, 742]
    titles = ["Выбор", "Навигация", "Сканирование", "Награда"]
    desc = [
        "Тематические маршруты по интересам",
        "Карта, путь и удобные остановки",
        "QR, факты, квизы и аудио",
        "Баллы, итоговый тест и бонусы",
    ]
    for pos, title, body, number in zip(x_positions, titles, desc, range(1, 5)):
        c.setFillColor(colors.Color(1, 1, 1, alpha=0.07))
        c.roundRect(pos, 180, 170, 176, 24, fill=1, stroke=0)
        c.setFillColor(BLUE)
        c.circle(pos + 28, 326, 15, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont(FONT_BOLD, 13)
        c.drawCentredString(pos + 28, 321, str(number))
        c.setFillColor(TEXT)
        c.setFont(FONT_BOLD, 18)
        c.drawString(pos + 20, 282, title)
        draw_text_block(c, wrap_lines(body, 18), pos + 20, 252, size=13, leading=17, color=MUTED)
        if pos < x_positions[-1]:
            c.setStrokeColor(BLUE_SOFT)
            c.setLineWidth(2)
            c.line(pos + 170, 266, pos + 208, 266)
            c.line(pos + 202, 272, pos + 208, 266)
            c.line(pos + 202, 260, pos + 208, 266)
    draw_footer(c, index)


def draw_metrics_slide(c: canvas.Canvas, slide: dict, index: int) -> None:
    c.setFillColor(BG)
    c.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 30)
    c.drawString(42, 454, slide["title"])
    c.setFont(FONT_REGULAR, 18)
    c.drawString(42, 426, slide["subtitle"])
    c.setFillColor(MUTED)
    c.setFont(FONT_REGULAR, 14)
    c.drawString(42, 482, slide["tag"])

    metrics = [
        ("4", "маршрута"),
        ("11", "точек"),
        ("12", "остановок"),
        ("2", "награды"),
    ]
    x_positions = [52, 248, 444, 640]
    for pos, (value, label) in zip(x_positions, metrics):
        c.setFillColor(colors.Color(1, 1, 1, alpha=0.07))
        c.roundRect(pos, 310, 160, 118, 24, fill=1, stroke=0)
        c.setFillColor(BLUE)
        c.setFont(FONT_BOLD, 34)
        c.drawCentredString(pos + 80, 368, value)
        c.setFillColor(TEXT)
        c.setFont(FONT_REGULAR, 16)
        c.drawCentredString(pos + 80, 338, label)

    draw_text_block(c, [slide["bullets"][1]], 58, 238, size=17, leading=22, color=TEXT)
    draw_text_block(c, [slide["bullets"][2]], 58, 200, size=17, leading=22, color=TEXT)
    draw_text_block(c, [slide["bullets"][3]], 58, 162, size=17, leading=22, color=TEXT)

    c.setFillColor(PANEL_ALT)
    c.roundRect(660, 110, 236, 150, 24, fill=1, stroke=0)
    image = ImageReader(str(IMAGES["sunset"]))
    c.drawImage(image, 678, 128, width=200, height=114, preserveAspectRatio=True, mask="auto")
    draw_footer(c, index)


def draw_partner_slide(c: canvas.Canvas, slide: dict, index: int) -> None:
    c.setFillColor(BG)
    c.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 30)
    c.drawString(42, 454, slide["title"])
    c.setFont(FONT_REGULAR, 18)
    c.drawString(42, 426, slide["subtitle"])
    c.setFillColor(MUTED)
    c.setFont(FONT_REGULAR, 14)
    c.drawString(42, 482, slide["tag"])
    columns = [
        ("Город", slide["bullets"][0], BLUE),
        ("Музеи", slide["bullets"][1], GREEN),
        ("Партнеры", slide["bullets"][2], RED),
    ]
    x_positions = [54, 330, 606]
    for pos, (title, body, accent) in zip(x_positions, columns):
        c.setFillColor(colors.Color(1, 1, 1, alpha=0.06))
        c.roundRect(pos, 156, 244, 220, 26, fill=1, stroke=0)
        c.setFillColor(accent)
        c.rect(pos, 156, 10, 220, fill=1, stroke=0)
        c.setFillColor(TEXT)
        c.setFont(FONT_BOLD, 20)
        c.drawString(pos + 28, 336, title)
        draw_text_block(c, wrap_lines(body, 24), pos + 28, 300, size=14, leading=18, color=MUTED)
    draw_text_block(c, wrap_lines(slide["bullets"][3], 92), 54, 116, size=14, leading=18, color=TEXT, bold=True)
    draw_footer(c, index)


def draw_architecture_slide(c: canvas.Canvas, slide: dict, index: int) -> None:
    c.setFillColor(BG)
    c.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 30)
    c.drawString(42, 454, slide["title"])
    c.setFont(FONT_REGULAR, 18)
    c.drawString(42, 426, slide["subtitle"])
    c.setFillColor(MUTED)
    c.setFont(FONT_REGULAR, 14)
    c.drawString(42, 482, slide["tag"])

    blocks = [
        (60, 238, 196, 84, "Frontend", "React + Tailwind\nPWA / mobile UI"),
        (384, 238, 196, 84, "Backend", "Node.js + Express\nAPI / progress / rewards"),
        (708, 238, 196, 84, "Интеграции", "2GIS / QR / offline\nконтент и маршруты"),
    ]
    for x, y, w, h, title, body in blocks:
        c.setFillColor(PANEL_ALT)
        c.roundRect(x, y, w, h, 18, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#111827"))
        c.setFont(FONT_BOLD, 16)
        c.drawString(x + 18, y + 54, title)
        draw_text_block(c, body.splitlines(), x + 18, y + 34, size=12, leading=15, color=colors.HexColor("#344054"))
    c.setStrokeColor(BLUE)
    c.setLineWidth(3)
    c.line(256, 280, 384, 280)
    c.line(580, 280, 708, 280)
    c.setFillColor(BLUE)
    c.circle(320, 280, 5, fill=1, stroke=0)
    c.circle(644, 280, 5, fill=1, stroke=0)

    lower_blocks = [
        (118, 118, 210, 72, "Контент", "Маршруты, точки, факты, аудио"),
        (374, 104, 210, 86, "Пользовательский прогресс", "Сканы, баллы, тесты,\nизбранное и профиль"),
        (630, 118, 210, 72, "Награды и партнеры", "Промокоды, магазин,\nбудущие платные сценарии"),
    ]
    for x, y, w, h, title, body in lower_blocks:
        c.setFillColor(colors.Color(1, 1, 1, alpha=0.08))
        c.roundRect(x, y, w, h, 18, fill=1, stroke=0)
        c.setFillColor(TEXT)
        c.setFont(FONT_BOLD, 15)
        c.drawString(x + 16, y + h - 22, title)
        draw_text_block(c, body.splitlines(), x + 16, y + h - 42, size=12, leading=15, color=MUTED)

    c.setStrokeColor(BLUE_SOFT)
    c.setLineWidth(1.5)
    c.line(482, 238, 482, 190)
    c.line(222, 190, 742, 190)
    c.line(222, 190, 222, 168)
    c.line(482, 190, 482, 154)
    c.line(742, 190, 742, 168)
    draw_footer(c, index)


def draw_roadmap_slide(c: canvas.Canvas, slide: dict, index: int) -> None:
    c.setFillColor(BG)
    c.rect(0, 0, PAGE_WIDTH, PAGE_HEIGHT, fill=1, stroke=0)
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 30)
    c.drawString(42, 454, slide["title"])
    c.setFont(FONT_REGULAR, 18)
    c.drawString(42, 426, slide["subtitle"])
    c.setFillColor(MUTED)
    c.setFont(FONT_REGULAR, 14)
    c.drawString(42, 482, slide["tag"])

    stages = [
        ("Этап 1", slide["bullets"][0]),
        ("Этап 2", slide["bullets"][1]),
        ("Этап 3", slide["bullets"][2]),
    ]
    x_positions = [54, 338, 622]
    for pos, (title, body) in zip(x_positions, stages):
        c.setFillColor(colors.Color(1, 1, 1, alpha=0.06))
        c.roundRect(pos, 186, 244, 160, 26, fill=1, stroke=0)
        c.setFillColor(BLUE)
        c.roundRect(pos + 18, 300, 84, 28, 14, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont(FONT_BOLD, 14)
        c.drawCentredString(pos + 60, 308, title)
        draw_text_block(c, wrap_lines(body, 24), pos + 18, 274, size=14, leading=18, color=TEXT)
    c.setFillColor(colors.Color(1, 1, 1, alpha=0.09))
    c.roundRect(54, 80, 812, 70, 20, fill=1, stroke=0)
    draw_text_block(c, wrap_lines(slide["bullets"][3], 82), 74, 110, size=16, leading=20, color=TEXT, bold=True)
    draw_footer(c, index)


def build_slide_pdf() -> None:
    c = canvas.Canvas(str(SLIDE_PDF), pagesize=landscape((PAGE_HEIGHT, PAGE_WIDTH)))
    c.setPageSize((PAGE_WIDTH, PAGE_HEIGHT))

    draw_cover(c, SLIDES[0], 0)
    c.showPage()
    draw_standard_slide(c, SLIDES[1], 1, image_key="kremlin", panel_style="right")
    c.showPage()
    draw_standard_slide(c, SLIDES[2], 2, image_key="pryanik", panel_style="right")
    c.showPage()
    draw_process_slide(c, SLIDES[3], 3)
    c.showPage()
    draw_metrics_slide(c, SLIDES[4], 4)
    c.showPage()
    draw_standard_slide(c, SLIDES[5], 5, image_key="vdv", panel_style="right")
    c.showPage()
    draw_partner_slide(c, SLIDES[6], 6)
    c.showPage()
    draw_architecture_slide(c, SLIDES[7], 7)
    c.showPage()
    draw_roadmap_slide(c, SLIDES[8], 8)
    c.showPage()
    c.save()


def build_notes_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("На Рязань — текст для выступающих", level=0)
    title.runs[0].font.name = "Arial"

    intro = doc.add_paragraph(
        "Документ содержит готовые тезисы для устного выступления по каждому слайду. "
        "Текст написан в живой форме, чтобы им можно было пользоваться как сценарием или как основой для собственной подачи."
    )
    intro.style = doc.styles["Normal"]

    for idx, slide in enumerate(SLIDES, start=1):
        doc.add_heading(f"Слайд {idx}. {slide['title']}", level=1)
        doc.add_paragraph(slide["subtitle"])
        bullet_title = doc.add_paragraph()
        bullet_title.add_run("Что на слайде: ").bold = True
        bullet_title.add_run("; ".join(slide["bullets"]))
        for note in slide["notes"]:
            doc.add_paragraph(note, style="List Bullet")

    doc.save(str(NOTES_DOCX))


def build_notes_pdf() -> None:
    c = canvas.Canvas(str(NOTES_PDF), pagesize=(595, 842))
    width, height = 595, 842
    margin_x = 48
    y = height - 56

    def new_page() -> None:
        nonlocal y
        c.showPage()
        y = height - 56

    def ensure_space(lines_needed: int) -> None:
        nonlocal y
        if y - (lines_needed * 18) < 56:
            new_page()

    c.setFillColor(BG)
    c.rect(0, 0, width, height, fill=1, stroke=0)
    c.setFillColor(TEXT)
    c.setFont(FONT_BOLD, 22)
    c.drawString(margin_x, y, "На Рязань — текст для выступающих")
    y -= 34
    draw_text_block(
        c,
        wrap_lines(
            "Сценарий подготовлен по слайдам презентации. Каждый блок можно читать почти дословно или использовать как опорные тезисы.",
            68,
        ),
        margin_x,
        y,
        size=12,
        leading=16,
        color=MUTED,
    )
    y -= 42

    for idx, slide in enumerate(SLIDES, start=1):
        needed = 8 + sum(max(1, len(wrap_lines(note, 76))) for note in slide["notes"])
        ensure_space(needed)
        if y < height - 56:
            c.setFillColor(BG)
            c.rect(0, 0, width, height, fill=1, stroke=0)
        c.setFillColor(BLUE)
        c.roundRect(margin_x, y - 4, 92, 20, 10, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont(FONT_BOLD, 11)
        c.drawCentredString(margin_x + 46, y + 2, f"Слайд {idx}")
        y -= 26
        c.setFillColor(TEXT)
        c.setFont(FONT_BOLD, 16)
        c.drawString(margin_x, y, slide["title"])
        y -= 18
        draw_text_block(c, wrap_lines(slide["subtitle"], 70), margin_x, y, size=12, leading=15, color=MUTED)
        y -= 28
        for note in slide["notes"]:
            wrapped = wrap_lines(note, 76)
            c.setFillColor(BLUE_SOFT)
            c.circle(margin_x + 4, y + 5, 2.7, fill=1, stroke=0)
            draw_text_block(c, wrapped, margin_x + 14, y, size=12, leading=16, color=TEXT)
            y -= 16 * len(wrapped) + 10
        y -= 10
    c.save()


def build_outline_md() -> None:
    lines = [
        "# На Рязань — структура презентации",
        "",
        "Figma Slides generation in this session was blocked by workspace deck-role restrictions, so this local package is the production-ready fallback.",
        "",
    ]
    for idx, slide in enumerate(SLIDES, start=1):
        lines.append(f"## Слайд {idx}. {slide['title']}")
        lines.append(f"Подзаголовок: {slide['subtitle']}")
        lines.append("")
        for bullet in slide["bullets"]:
            lines.append(f"- {bullet}")
        lines.append("")
    OUTLINE_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    register_fonts()
    build_slide_pdf()
    build_notes_docx()
    build_notes_pdf()
    build_outline_md()
    print(f"Created: {SLIDE_PDF}")
    print(f"Created: {NOTES_DOCX}")
    print(f"Created: {NOTES_PDF}")
    print(f"Created: {OUTLINE_MD}")


if __name__ == "__main__":
    main()
